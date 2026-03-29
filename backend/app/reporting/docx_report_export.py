"""Minimal generated DOCX report from a structured analysis payload (Stage 15.3).

Accepts the same JSON-shaped dict as the upload success response. Issue extraction
mirrors ``issuesFromUploadResponse`` / ``issueFragmentPlainText`` in ``upload.js``.
"""

from __future__ import annotations

import re
from io import BytesIO
from pathlib import Path
from typing import Any, Mapping
from urllib.parse import quote

from docx import Document

# Russian labels aligned with ``SEVERITY_LABEL_RU`` / ``CHECK_KEY_LABEL_RU`` in upload.js.
_SEVERITY_LABEL_RU: dict[str, str] = {
    "critical": "Критично",
    "warning": "Предупреждение",
    "recommendation": "Рекомендация",
}

_CHECK_KEY_LABEL_RU: dict[str, str] = {
    "vague_wording_check": "Неконкретные формулировки",
    "unverifiable_requirements_check": "Проверяемость требований",
    "figure_references_check": "Ссылки на рисунки",
    "table_references_check": "Ссылки на таблицы",
    "appendix_references_check": "Ссылки на приложения",
    "terminology_consistency_check": "Согласованность терминов",
    "duplicate_formulations_check": "Повторы формулировок",
}

_WIN_RESERVED = re.compile(r"^(CON|PRN|AUX|NUL|COM[1-9]|LPT[1-9])(\.|$)", re.IGNORECASE)


def safe_export_filename_stem(original_filename: str | None) -> str:
    """Match ``safeExportFilenameStem`` in ``static/js/upload.js`` (safe client stem)."""
    raw = (original_filename or "").strip() or "document"
    base = Path(raw).name
    dot = base.rfind(".")
    stem = base[:dot] if dot > 0 else base
    cleaned = re.sub(r'[\x00-\x1f<>:"/\\|?*]+', "_", stem)
    cleaned = _WIN_RESERVED.sub(r"_\1", cleaned).strip()
    cleaned = (cleaned[:120] if cleaned else "") or "document"
    return cleaned


def content_disposition_attachment(filename: str) -> str:
    """RFC 5987 ``Content-Disposition`` with ASCII fallback and UTF-8 filename*."""
    ascii_name = filename.encode("ascii", "replace").decode("ascii").replace("?", "_")
    ascii_name = ascii_name.replace('"', "_").replace("\\", "_")
    if not ascii_name.strip():
        ascii_name = "report.docx"
    enc = quote(filename, safe="")
    return f'attachment; filename="{ascii_name}"; filename*=UTF-8\'\'{enc}'


def issues_from_analysis_payload(data: Mapping[str, Any]) -> list[dict[str, Any]]:
    raw = data.get("issues")
    if isinstance(raw, list):
        return [x for x in raw if isinstance(x, dict)]
    rep = data.get("report")
    if isinstance(rep, dict):
        inner = rep.get("issues")
        if isinstance(inner, list):
            return [x for x in inner if isinstance(x, dict)]
    return []


def conclusion_from_payload(data: Mapping[str, Any]) -> str:
    rep = data.get("report")
    if not isinstance(rep, dict):
        return ""
    summ = rep.get("summary")
    if not isinstance(summ, dict):
        return ""
    c = summ.get("conclusion")
    if c is None:
        return ""
    s = str(c).strip()
    return s


def _nonempty_str(value: object | None) -> str:
    if value is None:
        return ""
    s = str(value).strip()
    return s


def issue_fragment_plain_text(issue: Mapping[str, Any]) -> str:
    s = _nonempty_str(issue.get("fragment_text"))
    if s:
        return s
    s = _nonempty_str(issue.get("excerpt"))
    if s:
        return s
    meta = issue.get("metadata")
    if isinstance(meta, dict):
        src = meta.get("source")
        if isinstance(src, dict):
            s = _nonempty_str(src.get("text_excerpt"))
            if s:
                return s
        s = _nonempty_str(meta.get("text_excerpt"))
        if s:
            return s
    return ""


def _add_labeled_paragraph(doc: Document, label: str, body: str) -> None:
    p = doc.add_paragraph()
    run = p.add_run(f"{label} ")
    run.bold = True
    p.add_run(body)


def build_analysis_docx_bytes(payload: Mapping[str, Any]) -> tuple[bytes, str]:
    """Return ``(docx_bytes, suggested_filename)`` for the given analysis dict."""
    issues = issues_from_analysis_payload(payload)
    conclusion = conclusion_from_payload(payload)
    name_hint = _nonempty_str(payload.get("filename"))
    out_name = f"{safe_export_filename_stem(name_hint)}-report.docx"

    doc = Document()
    doc.add_heading("Отчёт о проверке технического задания", level=1)

    _add_labeled_paragraph(doc, "Исходный файл:", name_hint or "—")

    if conclusion:
        doc.add_paragraph(conclusion)
    else:
        doc.add_paragraph("Краткий вывод в данных отсутствует.")

    doc.add_paragraph(f"Число замечаний: {len(issues)}")

    doc.add_heading("Список замечаний", level=2)

    if not issues:
        doc.add_paragraph("Замечаний нет.")
    else:
        for i, issue in enumerate(issues, start=1):
            doc.add_heading(f"Замечание {i}", level=3)
            sev_key = _nonempty_str(issue.get("severity"))
            sev_ru = _SEVERITY_LABEL_RU.get(sev_key, sev_key or "")
            if sev_ru:
                _add_labeled_paragraph(doc, "Серьёзность:", sev_ru)
            ck = _nonempty_str(issue.get("check_key"))
            if ck:
                ck_ru = _CHECK_KEY_LABEL_RU.get(ck, ck)
                _add_labeled_paragraph(doc, "Тип проверки:", ck_ru)
            msg = _nonempty_str(issue.get("message"))
            if msg:
                _add_labeled_paragraph(doc, "Описание:", msg)
            st = _nonempty_str(issue.get("section_title"))
            if st:
                _add_labeled_paragraph(doc, "Раздел:", st)
            frag = issue_fragment_plain_text(issue)
            if frag:
                _add_labeled_paragraph(doc, "Фрагмент:", frag)
            rec = _nonempty_str(issue.get("recommendation"))
            if rec:
                _add_labeled_paragraph(doc, "Рекомендация:", rec)

    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue(), out_name
