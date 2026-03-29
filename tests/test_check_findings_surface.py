"""Stage 17.3: verify implemented checks produce real findings in issues and report."""

from __future__ import annotations

from io import BytesIO

from docx import Document
from fastapi.testclient import TestClient

from backend.app.checks.required_sections_check import check_required_sections_presence
from backend.app.checks.section_order_check import check_section_order
from backend.app.structure_builder import add_section_tree
from backend.app.main import app
from backend.app.reporting.issues_builder import build_document_issues_serialized
from backend.app.reporting.report_builder import build_analysis_report
from backend.app.structure_detector import detect_structure
from backend.app.structure_enricher import enrich_structure
from backend.app.rules.section_rules import REQUIRED_SECTIONS


def _enriched_from_blocks(blocks: list[str]) -> dict:
    raw = add_section_tree(detect_structure(blocks))
    return enrich_structure(raw, blocks)


def test_required_sections_missing_surface_as_issues() -> None:
    checks = check_required_sections_presence({"sections": []})
    issues = build_document_issues_serialized(checks)
    assert len(issues) == len(REQUIRED_SECTIONS)
    codes = {row["issue_code"] for row in issues}
    assert all(c.startswith("required_sections.missing.") for c in codes)
    assert {row["check_key"] for row in issues} == {"required_sections_check"}
    rep = build_analysis_report(issues)
    assert rep["summary"]["total_issues"] == len(REQUIRED_SECTIONS)
    assert rep["summary"]["status"] == "needs_review"
    assert rep["summary"]["by_check"]["required_sections_check"] == len(REQUIRED_SECTIONS)


def test_section_order_violation_surfaces_in_issues_and_report() -> None:
    blocks = [
        "1. Назначение",
        "A.",
        "Введение",
    ]
    enriched = _enriched_from_blocks(blocks)
    order = check_section_order(enriched)
    issues = build_document_issues_serialized(order)
    assert len(issues) == 1
    assert issues[0]["check_key"] == "section_order_check"
    assert issues[0]["issue_code"] == "section_order.violation"
    assert "Назначение" in issues[0]["message"]
    assert issues[0]["severity"] == "warning"
    rep = build_analysis_report(issues)
    assert rep["summary"]["total_issues"] == 1
    assert rep["summary"]["conclusion"].startswith("В документе обнаружены")


def test_content_presence_absent_surfaces_warning_issues() -> None:
    checks = {
        "purpose_check": {
            "is_present": False,
            "match_type": None,
            "matched_signal": None,
            "source": None,
        },
        "scope_check": {
            "is_present": False,
            "match_type": None,
            "matched_signal": None,
            "source": None,
        },
    }
    issues = build_document_issues_serialized(checks)
    codes = {row["issue_code"] for row in issues}
    assert "content.purpose_missing" in codes
    assert "content.scope_missing" in codes
    assert all(row["severity"] == "warning" for row in issues)


def test_content_presence_when_present_does_not_emit_issue() -> None:
    checks = {
        "purpose_check": {
            "is_present": True,
            "match_type": "section_title",
            "matched_signal": "Назначение",
            "source": {"section_title": "Назначение"},
        }
    }
    issues = build_document_issues_serialized(checks)
    assert not any(row["check_key"] == "purpose_check" for row in issues)


def test_numerical_characteristics_gap_surfaces_recommendation() -> None:
    checks = {
        "numerical_characteristics_check": {
            "has_numeric_findings": False,
            "item_count": 1,
            "items": [
                {
                    "text": "Система должна поддерживать отказоустойчивость.",
                    "has_numeric_characteristics": False,
                    "matched_signals": [],
                    "fragment_index": 2,
                    "section_title": "Требования",
                }
            ],
        }
    }
    issues = build_document_issues_serialized(checks)
    assert len(issues) == 1
    assert issues[0]["issue_code"] == "numerical.missing_in_obligation_fragment"
    assert issues[0]["severity"] == "recommendation"
    assert issues[0]["locations"] == [{"fragment_index": 2}]


def test_numerical_all_fragments_numeric_emits_no_issues() -> None:
    checks = {
        "numerical_characteristics_check": {
            "has_numeric_findings": True,
            "item_count": 1,
            "items": [
                {
                    "text": "Время отклика не более 2 с.",
                    "has_numeric_characteristics": True,
                    "matched_signals": ["time"],
                    "fragment_index": 0,
                }
            ],
        }
    }
    assert build_document_issues_serialized(checks) == []


def test_measurement_units_review_surfaces_when_unlinked() -> None:
    checks = {
        "measurement_units_check": {
            "has_findings": True,
            "item_count": 1,
            "items": [
                {
                    "text": "Допускается до 5 % отклонения.",
                    "has_numeric_value": True,
                    "has_unit": True,
                    "number_unit_linked": False,
                    "matched_units": ["%"],
                    "reason": "число и единица измерения присутствуют, но связь между ними неочевидна",
                    "fragment_index": 1,
                }
            ],
        }
    }
    issues = build_document_issues_serialized(checks)
    assert len(issues) == 1
    assert issues[0]["check_key"] == "measurement_units_check"
    assert issues[0]["issue_code"] == "measurement.unit_linkage_review"
    assert "связь" in issues[0]["message"].lower()


def test_measurement_linked_number_and_unit_emits_no_issue() -> None:
    checks = {
        "measurement_units_check": {
            "has_findings": False,
            "item_count": 1,
            "items": [
                {
                    "text": "Не более 10 секунд.",
                    "has_numeric_value": True,
                    "has_unit": True,
                    "number_unit_linked": True,
                    "matched_units": ["секунд"],
                    "reason": "числовое значение связано с единицей измерения",
                    "fragment_index": 0,
                }
            ],
        }
    }
    assert build_document_issues_serialized(checks) == []


def test_table_references_missing_declaration_surfaces() -> None:
    checks = {
        "table_references_check": {
            "missing_declarations": [3],
            "unreferenced_tables": [],
            "duplicate_declarations": [],
            "caption_numbering_gaps": [],
            "mentions": [
                {
                    "number": 3,
                    "block_index": 0,
                    "line_index": 0,
                    "mention_text": "табл. 3",
                    "line": "см. табл. 3",
                }
            ],
            "declarations": [],
        }
    }
    issues = build_document_issues_serialized(checks)
    assert len(issues) == 1
    assert issues[0]["issue_code"] == "table.missing_declaration"
    assert issues[0]["locations"] == [{"block_index": 0, "line_index": 0}]


def test_appendix_missing_declaration_surfaces() -> None:
    checks = {
        "appendix_references_check": {
            "mentioned_labels": ["А"],
            "declared_labels": [],
            "missing_declarations": ["А"],
            "unreferenced_appendices": [],
            "duplicate_declarations": [],
            "cyrillic_letter_sequence_gaps": [],
            "numeric_appendix_gaps": [],
            "has_issues": True,
            "mentions": [
                {
                    "label": "А",
                    "block_index": 1,
                    "line_index": 0,
                    "mention_text": "приложение А",
                    "line": "см. приложение А",
                }
            ],
            "declarations": [],
        }
    }
    issues = build_document_issues_serialized(checks)
    assert len(issues) == 1
    assert issues[0]["issue_code"] == "appendix.missing_declaration.А"
    assert issues[0]["locations"] == [{"block_index": 1, "line_index": 0}]


def test_vague_wording_empty_findings_emits_no_issues() -> None:
    checks = {
        "vague_wording_check": {
            "has_findings": False,
            "finding_count": 0,
            "findings": [],
        }
    }
    assert build_document_issues_serialized(checks) == []


def test_upload_docx_swapped_sections_and_vague_text_surfaces_issues() -> None:
    doc = Document()
    doc.add_paragraph("1. Назначение")
    doc.add_paragraph("Введение")
    doc.add_paragraph("Система должна работать быстро.")
    buf = BytesIO()
    doc.save(buf)
    data = buf.getvalue()

    client = TestClient(app)
    response = client.post(
        "/api/upload",
        files={
            "file": (
                "spec.docx",
                data,
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )
        },
    )
    assert response.status_code == 200
    body = response.json()
    assert body["status"] == "accepted"
    issues = body["issues"]
    assert isinstance(issues, list)
    assert len(issues) >= 2
    check_keys = {i["check_key"] for i in issues}
    assert "section_order_check" in check_keys
    assert "vague_wording_check" in check_keys
    report = body["report"]
    assert report["summary"]["total_issues"] == len(issues)
    assert report["summary"]["status"] == "needs_review"
    order_block = body["checks"].get("section_order_check") or {}
    assert order_block.get("is_correct") is False
    vague_block = body["checks"].get("vague_wording_check") or {}
    assert vague_block.get("has_findings") is True
