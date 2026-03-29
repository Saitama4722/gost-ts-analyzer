"""Tests for DOCX text extraction."""

import zipfile
from io import BytesIO
from pathlib import Path

import pytest
from docx import Document
from fastapi.testclient import TestClient

from backend.app.docx_extraction import (
    DocxExtractionError,
    extract_docx_from_bytes,
    extract_docx_from_path,
)
from backend.app.main import UPLOADS_DIR, app


def _minimal_docx_bytes(*lines: str) -> bytes:
    doc = Document()
    for line in lines:
        doc.add_paragraph(line)
    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()


def test_extract_docx_from_bytes_preserves_order_and_full_text() -> None:
    data = _minimal_docx_bytes("Первый", "", "Третий")
    result = extract_docx_from_bytes(data)
    assert result.paragraphs == ["Первый", "", "Третий"]
    assert result.full_text == "Первый\n\nТретий"


def test_extract_docx_from_bytes_strips_whitespace_per_paragraph() -> None:
    doc = Document()
    doc.add_paragraph("  trimmed  ")
    buf = BytesIO()
    doc.save(buf)
    result = extract_docx_from_bytes(buf.getvalue())
    assert result.paragraphs == ["trimmed"]


def test_extract_docx_empty_file_raises() -> None:
    with pytest.raises(DocxExtractionError, match="пуст"):
        extract_docx_from_bytes(b"")


def test_extract_docx_invalid_bytes_raises() -> None:
    with pytest.raises(DocxExtractionError, match="корректным"):
        extract_docx_from_bytes(b"not a docx")


def test_extract_docx_valid_zip_without_office_package_raises() -> None:
    buf = BytesIO()
    with zipfile.ZipFile(buf, "w", zipfile.ZIP_DEFLATED) as zf:
        zf.writestr("note.txt", "not a Word document")
    with pytest.raises(DocxExtractionError, match="корректным"):
        extract_docx_from_bytes(buf.getvalue())


def test_extract_docx_from_path_missing_raises(tmp_path: Path) -> None:
    missing = tmp_path / "nope.docx"
    with pytest.raises(DocxExtractionError, match="найден"):
        extract_docx_from_path(missing)


def test_upload_docx_includes_extraction() -> None:
    data = _minimal_docx_bytes("A", "B")
    client = TestClient(app)
    response = client.post(
        "/api/upload",
        files={
            "file": (
                "spec.docx",
                data,
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )
        },
    )
    assert response.status_code == 200
    body = response.json()
    assert body["status"] == "accepted"
    assert body["extraction"]["blocks"] == ["A", "B"]
    assert body["extraction"]["full_text"] == "A\nB"
    assert "checks" in body
    assert "issues" in body
    assert "report" in body
    assert isinstance(body["issues"], list)
    stored = body["stored_as"]
    assert (UPLOADS_DIR / stored).is_file()


def test_upload_docx_corrupt_returns_422() -> None:
    client = TestClient(app)
    response = client.post(
        "/api/upload",
        files={
            "file": (
                "bad.docx",
                b"PK\x03\x04broken",
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )
        },
    )
    assert response.status_code == 422
    payload = response.json()
    assert payload["status"] == "error"
    assert "корректным" in payload.get("message", "")


def test_upload_docx_plain_text_disguised_as_docx_returns_422() -> None:
    client = TestClient(app)
    response = client.post(
        "/api/upload",
        files={
            "file": (
                "report.docx",
                b"This is not a zip archive.",
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )
        },
    )
    assert response.status_code == 422
    assert response.json()["status"] == "error"


def test_upload_docx_valid_zip_not_word_returns_422() -> None:
    buf = BytesIO()
    with zipfile.ZipFile(buf, "w", zipfile.ZIP_DEFLATED) as zf:
        zf.writestr("data.txt", "x")
    client = TestClient(app)
    response = client.post(
        "/api/upload",
        files={
            "file": (
                "fake.docx",
                buf.getvalue(),
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )
        },
    )
    assert response.status_code == 422
    assert response.json()["status"] == "error"


def test_upload_unsupported_extension_returns_415() -> None:
    client = TestClient(app)
    response = client.post(
        "/api/upload",
        files={"file": ("notes.txt", b"hello", "text/plain")},
    )
    assert response.status_code == 415
    assert response.json()["status"] == "error"


def test_upload_two_docx_files_sequentially_independent() -> None:
    client = TestClient(app)
    first = _minimal_docx_bytes("Первый")
    second = _minimal_docx_bytes("Второй")
    r1 = client.post(
        "/api/upload",
        files={
            "file": (
                "one.docx",
                first,
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )
        },
    )
    r2 = client.post(
        "/api/upload",
        files={
            "file": (
                "two.docx",
                second,
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )
        },
    )
    assert r1.status_code == 200
    assert r2.status_code == 200
    b1 = r1.json()
    b2 = r2.json()
    assert b1["stored_as"] != b2["stored_as"]
    assert b1["extraction"]["full_text"] == "Первый"
    assert b2["extraction"]["full_text"] == "Второй"
    assert (UPLOADS_DIR / b1["stored_as"]).is_file()
    assert (UPLOADS_DIR / b2["stored_as"]).is_file()


def test_upload_docx_nearly_empty_accepted() -> None:
    doc = Document()
    buf = BytesIO()
    doc.save(buf)
    client = TestClient(app)
    response = client.post(
        "/api/upload",
        files={
            "file": (
                "emptyish.docx",
                buf.getvalue(),
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )
        },
    )
    assert response.status_code == 200
    body = response.json()
    assert body["status"] == "accepted"
    assert "extraction" in body
    assert body["extraction"]["full_text"] == ""
    assert body["extraction"]["blocks"] in ([], [""])
